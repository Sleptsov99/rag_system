"""
Document Processor
==================
Three responsibilities:
  1. DocumentLoader   — reads PDF / DOCX / TXT files into plain text
  2. TextPreprocessor — tokenisation, stop-word removal, lemmatisation
  3. TextChunker      — splits long texts into overlapping chunks
"""

from __future__ import annotations

import re
import string
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import nltk

# Download required NLTK data on first run (silent after that)
for _pkg in ("punkt", "punkt_tab", "stopwords"):
    try:
        nltk.data.find(f"tokenizers/{_pkg}" if _pkg.startswith("punkt") else f"corpora/{_pkg}")
    except LookupError:
        nltk.download(_pkg, quiet=True)

from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Optional: Russian lemmatiser (pymorphy2). Falls back gracefully if absent.
try:
    import pymorphy3 as pymorphy2
    _morph = pymorphy2.MorphAnalyzer()
    _PYMORPHY_AVAILABLE = True
except ImportError:
    _PYMORPHY_AVAILABLE = False

# Optional: PDF support
try:
    import PyPDF2
    _PDF_AVAILABLE = True
except ImportError:
    _PDF_AVAILABLE = False

# Optional: DOCX support
try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


# ---------------------------------------------------------------------------
# Data models
# ---------------------------------------------------------------------------

@dataclass
class Document:
    """A single piece of text with its origin metadata."""
    text: str
    source: str                          # file path or URL
    chunk_id: int = 0
    metadata: dict = field(default_factory=dict)

    def __repr__(self) -> str:
        preview = self.text[:80].replace("\n", " ")
        return f"Document(source={self.source!r}, chunk_id={self.chunk_id}, text={preview!r}…)"


# ---------------------------------------------------------------------------
# 1. Document Loader
# ---------------------------------------------------------------------------

class DocumentLoader:
    """
    Loads plain text from:
      • .txt  — UTF-8 text files
      • .pdf  — PDF documents (requires PyPDF2)
      • .docx — Word documents (requires python-docx)
    """

    def load_file(self, path: str | Path) -> str:
        path = Path(path)
        if not path.exists():
            raise FileNotFoundError(path)

        ext = path.suffix.lower()
        if ext == ".txt":
            return self._load_txt(path)
        elif ext == ".pdf":
            return self._load_pdf(path)
        elif ext == ".docx":
            return self._load_docx(path)
        else:
            raise ValueError(f"Unsupported format: {ext}")

    def load_directory(self, directory: str | Path) -> list[Document]:
        """Load all supported files in a directory."""
        directory = Path(directory)
        documents: list[Document] = []
        supported = {".txt", ".pdf", ".docx"}

        for file_path in sorted(directory.iterdir()):
            if file_path.suffix.lower() not in supported:
                continue
            try:
                text = self.load_file(file_path)
                if text.strip():
                    documents.append(Document(text=text, source=str(file_path)))
                    print(f"  Loaded: {file_path.name} ({len(text)} chars)")
            except Exception as exc:
                print(f"  Warning: could not load {file_path.name}: {exc}")

        return documents

    # --- private helpers ---

    @staticmethod
    def _load_txt(path: Path) -> str:
        return path.read_text(encoding="utf-8", errors="replace")

    @staticmethod
    def _load_pdf(path: Path) -> str:
        if not _PDF_AVAILABLE:
            raise ImportError("PyPDF2 is required for PDF support: pip install PyPDF2")
        pages: list[str] = []
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
        return "\n".join(pages)

    @staticmethod
    def _load_docx(path: Path) -> str:
        if not _DOCX_AVAILABLE:
            raise ImportError("python-docx is required: pip install python-docx")
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ---------------------------------------------------------------------------
# 2. Text Preprocessor
# ---------------------------------------------------------------------------

class TextPreprocessor:
    """
    Cleans and normalises text for embedding / search.

    Pipeline:
      raw text
        → lowercase
        → remove extra whitespace / punctuation
        → tokenise (NLTK word_tokenize)
        → remove stop words
        → lemmatise (pymorphy2 for Russian; falls back to stem-only for EN)
    """

    def __init__(self, language: str = "russian"):
        self.language = language
        try:
            self._stop_words = set(stopwords.words(language))
        except OSError:
            self._stop_words = set()

    def clean(self, text: str) -> str:
        """Light cleaning: collapse whitespace, fix encoding artefacts."""
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"[^\w\s\-.,!?;:()\[\]{}'\"«»—–]", " ", text)
        return text.strip()

    def tokenize(self, text: str) -> list[str]:
        """Split text into word tokens (NLTK punkt tokeniser)."""
        return word_tokenize(text.lower(), language=self.language)

    def remove_stopwords(self, tokens: list[str]) -> list[str]:
        punct = set(string.punctuation + "«»—–")
        return [t for t in tokens if t not in self._stop_words and t not in punct]

    def lemmatize(self, tokens: list[str]) -> list[str]:
        """Return base forms. Uses pymorphy2 for Russian if available."""
        if _PYMORPHY_AVAILABLE:
            return [_morph.parse(t)[0].normal_form for t in tokens]
        # Fallback: return tokens as-is (no EN lemmatiser added to keep deps light)
        return tokens

    def preprocess(self, text: str) -> str:
        """Full pipeline → returns cleaned string (used for display / debug)."""
        text = self.clean(text)
        tokens = self.tokenize(text)
        tokens = self.remove_stopwords(tokens)
        tokens = self.lemmatize(tokens)
        return " ".join(tokens)

    def preprocess_for_embedding(self, text: str) -> str:
        """
        Minimal cleaning for embedding: keep sentence structure.
        Aggressive preprocessing *hurts* transformer embeddings.
        """
        return self.clean(text)


# ---------------------------------------------------------------------------
# 3. Text Chunker
# ---------------------------------------------------------------------------

class TextChunker:
    """
    Splits text into overlapping chunks.

    Uses a recursive strategy (similar to LangChain's
    RecursiveCharacterTextSplitter): tries to split on paragraph
    boundaries first, then sentence boundaries, then raw characters.
    """

    SEPARATORS = ["\n\n", "\n", ". ", "! ", "? ", " ", ""]

    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        if chunk_overlap >= chunk_size:
            raise ValueError("chunk_overlap must be smaller than chunk_size")
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def split_text(self, text: str) -> list[str]:
        """Return a list of text chunks."""
        return self._recursive_split(text, self.SEPARATORS)

    def split_documents(self, documents: list[Document]) -> list[Document]:
        """Chunk every document and return a flat list of chunk-Documents."""
        result: list[Document] = []
        for doc in documents:
            chunks = self.split_text(doc.text)
            for i, chunk in enumerate(chunks):
                result.append(Document(
                    text=chunk,
                    source=doc.source,
                    chunk_id=i,
                    metadata={**doc.metadata, "total_chunks": len(chunks)},
                ))
        return result

    # --- private ---

    def _recursive_split(self, text: str, separators: list[str]) -> list[str]:
        if not separators:
            return self._split_by_chars(text)

        sep, *rest = separators
        if sep == "":
            return self._split_by_chars(text)

        parts = text.split(sep)
        chunks: list[str] = []
        current = ""

        for part in parts:
            candidate = current + (sep if current else "") + part
            if len(candidate) <= self.chunk_size:
                current = candidate
            else:
                if current:
                    chunks.append(current)
                # part itself may be too long → recurse with next separator
                if len(part) > self.chunk_size:
                    chunks.extend(self._recursive_split(part, rest))
                    current = ""
                else:
                    current = part

        if current:
            chunks.append(current)

        return self._apply_overlap(chunks)

    def _split_by_chars(self, text: str) -> list[str]:
        chunks = []
        start = 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunks.append(text[start:end])
            start += self.chunk_size - self.chunk_overlap
        return chunks

    def _apply_overlap(self, chunks: list[str]) -> list[str]:
        """Re-attach overlap from previous chunk to the beginning of each chunk."""
        if self.chunk_overlap == 0 or len(chunks) <= 1:
            return chunks
        result = [chunks[0]]
        for i in range(1, len(chunks)):
            overlap_text = chunks[i - 1][-self.chunk_overlap:]
            result.append(overlap_text + chunks[i])
        return result
